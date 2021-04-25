import os, base64, datetime
from docx2pdf import convert
from docxtpl import DocxTemplate
from docx.shared import Inches


def getDateDict():

    now = datetime.datetime.now()
    day = int(now.day)
    month = int(now.month)
    if month == 1:
        month = "января"
    elif month == 2:
        month = "февраля"
    elif month == 3:
        month = "марта"
    elif month == 4:
        month = "апреля"
    elif month == 5:
        month = "мая"
    elif month == 6:
        month = "июня"
    elif month == 7:
        month = "июля"
    elif month == 8:
        month = "августа"
    elif month == 9:
        month = "сентября"
    elif month == 10:
        month = "октября"
    elif month == 11:
        month = "ноября"
    else:
        month = "декабря"
    year = str(now.year)
    year = int(year[2] + year[3])

    date = {
        "day": day,
        "month": month,
        "year": year
    }

    return date;


def fillAndConvert(info_folder, form_file_name, data, uuid, formData):
    doc = DocxTemplate(info_folder + form_file_name + ".docx")
    tables = doc.tables
    curFormData = formData.get(form_file_name)
    if curFormData is not None:
        table = int(curFormData.get("sign_table"))
        row = int(curFormData.get("sign_row"))
        cell = int(curFormData.get("sign_cell"))
        p = tables[table].rows[row].cells[cell].add_paragraph()
        r = p.add_run()
        r.add_picture(info_folder + "signature-user-" + uuid + '.png', width=Inches(3), height=Inches(.5))


    data.update(getDateDict())
    doc.render(data)

    doc.save(info_folder + form_file_name + "-user-" + uuid + ".docx")
    convert(info_folder + form_file_name + "-user-" + uuid + ".docx", info_folder + form_file_name + "-user-" + uuid + ".pdf")
    os.remove(info_folder + form_file_name + "-user-" + uuid + ".docx")


def decodeDataUrl(info_folder, code, uuid):
    with open(info_folder + "signature-user-" + uuid + '.png', 'wb') as f:
        f.write(base64.b64decode(code.split(',')[1].encode()))
