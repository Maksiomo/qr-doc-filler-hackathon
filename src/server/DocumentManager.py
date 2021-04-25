import os, base64
from docx2pdf import convert
from docxtpl import DocxTemplate
from docx.shared import Inches


def fillAndConvert(info_folder, form_file_name, data, uuid, formData):
    doc = DocxTemplate(info_folder + form_file_name + ".docx")
    tables = doc.tables
    curFormData = formData.get(form_file_name)
    if curFormData is not None:
        table = int(curFormData.get("sign_table"))
        row = int(curFormData.get("sign_row"))
        cell = int(curFormData.get("sign_cell"))
        p = tables[table].rows[row].cells[cell].add_paragraph()
        r = p.add_run()
        r.add_picture(info_folder + "signature-user-" + uuid + '.png', width=Inches(3), height=Inches(.5))

    doc.render(data)

    doc.save(info_folder + form_file_name + "-user-" + uuid + ".docx")
    convert(info_folder + form_file_name + "-user-" + uuid + ".docx", info_folder + form_file_name + "-user-" + uuid + ".pdf")
    os.remove(info_folder + form_file_name + "-user-" + uuid + ".docx")


def decodeDataUrl(info_folder, code, uuid):
    with open(info_folder + "signature-user-" + uuid + '.png', 'wb') as f:
        f.write(base64.b64decode(code.split(',')[1].encode()))
